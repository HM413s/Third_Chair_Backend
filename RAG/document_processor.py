from typing import List, Dict, Any
import os
from PIL import Image
import pytesseract
from sklearn.feature_extraction.text import TfidfVectorizer
from rank_bm25 import BM25Okapi
import numpy as np
import base64
from io import BytesIO
import PyPDF2
import docx
import chardet
import re

class DocumentProcessor:
    def __init__(self):
        self.vectorizer = TfidfVectorizer(
            max_features=4000,
            stop_words='english',
            ngram_range=(1, 2)
        )
        self.bm25 = None
        self.documents = []
    
    def _extract_document_structure(self, text: str, file_type: str = "text") -> Dict[str, Any]:
        """Extract document structure information for better location tracking"""
        structure = {
            "sections": [],
            "headers": [],
            "paragraphs": [],
            "lists": [],
            "tables": [],
            "total_lines": 0,
            "total_words": 0,
            "total_chars": len(text)
        }
        
        lines = text.split('\n')
        structure["total_lines"] = len(lines)
        structure["total_words"] = len(text.split())
        
        # Detect sections (various patterns)
        section_patterns = [
            r'^(Section|SECTION|Sec\.)\s+(\d+(?:\.\d+)*)\s*[:\-\.]?\s*(.*)$',
            r'^(Article|ARTICLE|Art\.)\s+([IVX]+|\d+)\s*[:\-\.]?\s*(.*)$',
            r'^(\d+\.\s+)(.+)$',
            r'^([A-Z]{2,}(?:\s+[A-Z]{2,})*)\s*$'  # ALL CAPS headers
        ]
        
        for line_num, line in enumerate(lines, 1):
            line_stripped = line.strip()
            if not line_stripped:
                continue
                
            # Check for section patterns
            for pattern in section_patterns:
                match = re.match(pattern, line_stripped)
                if match:
                    if len(match.groups()) >= 3:
                        section_type, section_num, section_title = match.groups()[:3]
                    elif len(match.groups()) == 2:
                        section_type, section_title = match.groups()
                        section_num = ""
                    else:
                        section_type = match.group(1)
                        section_num = ""
                        section_title = ""
                    
                    structure["sections"].append({
                        "type": section_type.strip(),
                        "number": section_num.strip(),
                        "title": section_title.strip(),
                        "line_number": line_num,
                        "char_start": sum(len(l) + 1 for l in lines[:line_num-1]),
                        "char_end": sum(len(l) + 1 for l in lines[:line_num]),
                        "full_text": line_stripped
                    })
                    break
            
            # Detect headers (lines that are significantly shorter and may be titles)
            if len(line_stripped) < 100 and len(line_stripped) > 5:
                # Check if it looks like a header (title case, all caps, or starts with number)
                if (line_stripped.istitle() or 
                    line_stripped.isupper() or 
                    re.match(r'^\d+\.', line_stripped)):
                    structure["headers"].append({
                        "text": line_stripped,
                        "line_number": line_num,
                        "char_start": sum(len(l) + 1 for l in lines[:line_num-1]),
                        "char_end": sum(len(l) + 1 for l in lines[:line_num]),
                        "type": "header"
                    })
            
            # Detect lists
            if re.match(r'^\s*[\-\*\+•]\s+', line_stripped) or re.match(r'^\s*\d+\.\s+', line_stripped):
                structure["lists"].append({
                    "text": line_stripped,
                    "line_number": line_num,
                    "char_start": sum(len(l) + 1 for l in lines[:line_num-1]),
                    "type": "list_item"
                })
        
        return structure

    def _find_containing_section(self, text: str, char_position: int, document_structure: Dict[str, Any]) -> Dict[str, Any]:
        """Find which section contains a given character position"""
        sections = document_structure.get("sections", [])
        containing_section = None
        
        for section in sections:
            if section["char_start"] <= char_position:
                if containing_section is None or section["char_start"] > containing_section["char_start"]:
                    containing_section = section
        
        return containing_section or {"type": "unknown", "number": "", "title": ""}

    def process_text(self, text: str, page_number: int = None, source_file: str = "unknown") -> List[Dict[str, Any]]:
        """Process text content and return chunks with enhanced metadata and location tracking"""
        # Extract document structure first
        document_structure = self._extract_document_structure(text)
        
        sentences = text.split('. ')
        chunks = [s.strip() + '.' for s in sentences if s.strip()]
        self.documents.extend(chunks)
        tfidf_matrix = self.vectorizer.fit_transform(chunks)
        # Always rebuild BM25 with the full corpus
        self.bm25 = BM25Okapi([doc.split() for doc in self.documents])
        processed = []
        char_idx = 0
        
        # Get text lines for line number calculation
        lines = text.split('\n')
        
        for i, chunk in enumerate(chunks):
            start_char = text.find(chunk, char_idx)
            end_char = start_char + len(chunk)
            char_idx = end_char
            
            # Calculate line number
            line_number = text[:start_char].count('\n') + 1
            
            # Find containing section
            containing_section = self._find_containing_section(text, start_char, document_structure)
            
            # Find nearby headers for context
            nearby_headers = []
            for header in document_structure.get("headers", []):
                if abs(header["line_number"] - line_number) <= 5:  # Within 5 lines
                    nearby_headers.append(header)
            
            embedding = tfidf_matrix[i].toarray()[0].tolist()
            # Ensure embedding is exactly 4000 dimensions
            if len(embedding) > 4000:
                embedding = embedding[:4000]
            elif len(embedding) < 4000:
                embedding = embedding + [0.0] * (4000 - len(embedding))
            
            # Enhanced metadata with comprehensive location tracking
            enhanced_metadata = {
                # Basic location info - ensure integers not None
                "page_number": page_number if page_number is not None else 0,
                "line_number": line_number,
                "start_char": start_char,
                "end_char": end_char,
                "sentence_id": i,
                "source_file": source_file,
                
                # Document structure context
                "containing_section": containing_section,
                "nearby_headers": nearby_headers[:3],  # Limit to 3 nearest headers
                "document_structure_summary": {
                    "total_sections": len(document_structure.get("sections", [])),
                    "total_headers": len(document_structure.get("headers", [])),
                    "total_lines": document_structure.get("total_lines", 0),
                    "total_words": document_structure.get("total_words", 0)
                },
                
                # Content analysis
                "word_count": len(chunk.split()),
                "char_count": len(chunk),
                "is_list_item": any(chunk.strip().startswith(prefix) for prefix in ['-', '*', '•', '1.', '2.', '3.']),
                "contains_reference": bool(re.search(r'(?:Section|Article|Clause|Exhibit)\s+\d+', chunk)),
                "contains_definition": bool(re.search(r'"[^"]+"\s+(?:means|shall mean|is defined as)', chunk)),
                
                # Location context for agents
                "location_context": {
                    "file_path": source_file,
                    "absolute_position": f"Line {line_number}, Char {start_char}-{end_char}",
                    "section_path": f"{containing_section.get('type', 'Unknown')} {containing_section.get('number', '')}".strip(),
                    "section_title": containing_section.get('title', ''),
                    "relative_position_in_section": start_char - containing_section.get('char_start', 0) if containing_section.get('char_start') else 0,
                    "surrounding_context": text[max(0, start_char-200):end_char+200],
                    "document_type": "text"
                }
            }
            
            processed.append({
                "content": chunk,
                "type": "text",
                "embedding": embedding,
                "metadata": enhanced_metadata
            })
        return processed
    
    def process_image(self, image_path: str, page_number: int = None, source_file: str = "unknown") -> Dict[str, Any]:
        """Process image and extract text using OCR, store enhanced location metadata"""
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            buffered = BytesIO()
            image.save(buffered, format="PNG")
            img_str = base64.b64encode(buffered.getvalue()).decode()
            tfidf_vector = self.vectorizer.fit_transform([text])
            self.documents.append(text)
            # Always rebuild BM25 with the full corpus
            self.bm25 = BM25Okapi([doc.split() for doc in self.documents])
            embedding = tfidf_vector.toarray()[0].tolist()
            if len(embedding) > 1000:
                embedding = embedding[:1000]
            elif len(embedding) < 1000:
                embedding = embedding + [0.0] * (1000 - len(embedding))
            
            # Enhanced image metadata
            enhanced_metadata = {
                "page_number": page_number,
                "source_file": source_file,
                "image_format": image.format,
                "image_size": image.size,
                "image_mode": image.mode,
                "ocr_text_length": len(text),
                "ocr_word_count": len(text.split()),
                "has_text_content": bool(text.strip()),
                
                # Location context for images
                "location_context": {
                    "file_path": source_file,
                    "media_type": "image",
                    "page_reference": f"Page {page_number}" if page_number else "Single image",
                    "content_type": "OCR extracted text",
                    "image_dimensions": f"{image.size[0]}x{image.size[1]}" if image.size else "unknown"
                }
            }
            
            return {
                "content": text,
                "type": "image",
                "image_base64": img_str,
                "embedding": embedding,
                "metadata": enhanced_metadata
            }
        except Exception as e:
            print(f"Error processing image {image_path}: {str(e)}")
            return None

    def detect_encoding(self, file_path: str) -> str:
        """Detect file encoding"""
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            return result['encoding'] or 'utf-8'

    def extract_text_from_pdf(self, file_path: str, source_file: str = None) -> List[Dict[str, Any]]:
        """Extract text from PDF file and return chunks with enhanced metadata"""
        text_chunks = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        page_chunks = self.process_text(text, page_number=page_num, source_file=source_file)
                        # Add PDF-specific metadata
                        for chunk in page_chunks:
                            chunk["metadata"]["pdf_page_count"] = len(pdf_reader.pages)
                            chunk["metadata"]["pdf_page_number"] = page_num
                            chunk["metadata"]["location_context"]["document_type"] = "PDF"
                            chunk["metadata"]["location_context"]["page_info"] = f"Page {page_num} of {len(pdf_reader.pages)}"
                        text_chunks.extend(page_chunks)
        except Exception as e:
            print(f"Error extracting text from PDF {file_path}: {str(e)}")
        return text_chunks

    def extract_text_from_docx(self, file_path: str, source_file: str = None) -> List[Dict[str, Any]]:
        """Extract text from DOCX file and return chunks with enhanced metadata"""
        chunks = []
        try:
            doc = docx.Document(file_path)
            full_text = ""
            paragraph_map = {}
            
            # Build full text and paragraph mapping
            for para_id, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                if para_text.strip():
                    start_pos = len(full_text)
                    full_text += para_text + "\n"
                    end_pos = len(full_text)
                    paragraph_map[para_id] = {
                        "text": para_text,
                        "start_char": start_pos,
                        "end_char": end_pos,
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    }
            
            # Process the full text
            text_chunks = self.process_text(full_text, source_file=source_file)
            
            # Enhance with DOCX-specific metadata
            for chunk in text_chunks:
                chunk_start = chunk["metadata"]["start_char"]
                
                # Find which paragraph this chunk belongs to
                containing_paragraph = None
                for para_id, para_info in paragraph_map.items():
                    if para_info["start_char"] <= chunk_start < para_info["end_char"]:
                        containing_paragraph = {
                            "paragraph_id": para_id,
                            "paragraph_style": para_info["style"],
                            "paragraph_text": para_info["text"][:100] + "..." if len(para_info["text"]) > 100 else para_info["text"]
                        }
                        break
                
                chunk["metadata"]["docx_paragraph"] = containing_paragraph
                chunk["metadata"]["docx_total_paragraphs"] = len(doc.paragraphs)
                chunk["metadata"]["location_context"]["document_type"] = "DOCX"
                chunk["metadata"]["location_context"]["paragraph_info"] = containing_paragraph
            
            chunks.extend(text_chunks)
                    
        except Exception as e:
            print(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return chunks

    def process_file(self, file_path: str, source_file: str = None) -> List[Dict[str, Any]]:
        """Process a file and return processed chunks with enhanced metadata"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            # Use provided source_file or extract from file_path
            file_name = source_file or os.path.basename(file_path)
            
            print(f"Processing file: {file_name}")
            
            if file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                result = self.process_image(file_path, source_file=file_name)
                return [result] if result else []
            elif file_extension == '.pdf':
                return self.extract_text_from_pdf(file_path, source_file=file_name)
            elif file_extension == '.docx':
                return self.extract_text_from_docx(file_path, source_file=file_name)
            else:
                # Plain text file
                encoding = self.detect_encoding(file_path)
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                return self.process_text(text, source_file=file_name)
            return []
        except Exception as e:
            print(f"Error processing file {file_path}: {str(e)}")
            return []
    
    def get_similar_documents(self, query: str, top_k: int = 5) -> List[int]:
        """Get indices of similar documents using BM25"""
        if self.bm25 is None:
            return []
        tokenized_query = query.split()
        scores = self.bm25.get_scores(tokenized_query)
        return np.argsort(scores)[-top_k:][::-1] 